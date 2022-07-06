import streamlit as st
import pandas as pd
from docxtpl import DocxTemplate
from num2words import num2words
from datetime import datetime
from docx import Document
from docxcompose.composer import Composer
import numpy as np
st.set_page_config(page_title='Dede Saputra App', page_icon = "D", layout = 'centered', initial_sidebar_state = 'auto')

st.title("Aplikasi Kuitansi")
st.sidebar.title("Menu")
option = st.sidebar.selectbox("Pilih jenis kuitansi", ("Transpor Lokal", "Kegiatan", "Perjadin"))
if option == 'Transpor Lokal':
    st.write("Ini menggunakan **format kuitansi transportasi lokal**, jika ingin membuat kuitansi lain pilih jenis kuitansi di Menu")
    form = st.form("template kuitansi")
    with form:
        col1, col2 = st.columns(2)
        nama = col1.text_input('Nama')
        nip = col2.text_input("NIP")
        layanan = col1.text_input("Layanan")
        mak = col2.text_input("MAK")
        tanggal = col1.date_input("Tanggal Kegiatan")
        kegiatan = col1.text_area("Nama Kegiatan")
        lokasi = col2.text_input("Lokasi Kegiatan")
        nilai = col2.text_input("Nominal")
        submit = form.form_submit_button("Kirim")
        tgl = tanggal.strftime("%d %B %Y")
        tgl2 = tgl.replace("January", "Januari").replace("February", "Februari").replace("March", "Maret").replace("May", "Mei").replace("June", "Juni").replace("July", "Juli").replace("August", "Agustus").replace("October", "Oktober").replace("December", "Desember")
    if submit:
        uang1 = float(nilai)
        thousands_separator = "."
        fractional_separator = ","
        currency = "Rp{:,.2f}".format(uang1)
        if thousands_separator == ".":
            main_currency, fractional_currency = currency.split(".")[0], currency.split(".")[1]
            new_main_currency = main_currency.replace(",", ".")
            currency = new_main_currency + fractional_separator + fractional_currency
        doc = DocxTemplate("kuitansiperjadinapp.docx")
        context = {
            "nama": nama,
            "nip": nip,
            "mak": mak,
            "kegiatan": kegiatan,
            "lokasi": lokasi,
            "tanggal": tgl2,
            "layanan": layanan,
            "terbilang": num2words(int(nilai), lang='id').title() + " Rupiah",
            "uang": currency
        }
        #st.write(context)
        output_name = f'download/{context["nama"]}.docx'
        doc.render(context)   
        doc.save(output_name)
        #convert(output_name, "hasil.pdf")
        with open(output_name, "rb") as file:
        #    btn = st.download_button(
         #           label="Download PDF",
          #          data=file,
           #         file_name="hasil.pdf",
            #        mime="application/octet-stream"
             #   )
            st.success("🎉 File kuitansi telah selesai dibuat")
            # st.write(html, unsafe_allow_html=True)
            # st.write("")
            st.download_button(
                "⬇️ Download File",
                data=file,
                file_name="Hasil.docx",
                mime="application/octet-stream",
            )


       
        
   
if option == 'Kegiatan':
    form = st.form("Upload file")
    with form:
        st.markdown(
        "**PERHATIAN: ** Untuk menggunakan aplikasi ini, silahkan unduh format nominatif[ di sini](https://docs.google.com/spreadsheets/d/10LDb3d-pbSRKuio8zY-HIS2kmP03ehhn/edit#gid=1302452135). Dimohon untuk mengisi nominatif sesuai dengan format file unduhan, tanpa nama kolom dan nama sheet. "
        )
        st.write("**Ini adalah format kuitansi perjalanan dinas ke daerah, jika ingin membuat kuitansi jenis lain silahkan klik menu**")

        #st.header("Proses pembuatan kuitansi")
        excelfile = st.file_uploader("Unggah File Nominatif")
        submit = st.form_submit_button("Proses Nominatif")
    if submit:
        try:
            # Convert Excel sheet to pandas dataframe
            df = pd.read_excel(excelfile, sheet_name="Sheet1")

            # Keep only date part YYYY-MM-DD (not the time)
            df["TODAY"] = pd.to_datetime(df["TODAY"]).dt.date
            df["TODAY_IN_ONE_WEEK"] = pd.to_datetime(df["TODAY_IN_ONE_WEEK"]).dt.date
            np_array = df["VENDOR"].to_numpy()
            #st.write(list(np_array))
            # Iterate over each row in df and render word document
            for record in df.to_dict(orient="records"):
                doc = DocxTemplate(f"pages/vendor-contract.docx")
                doc.render(record)
                output_path = f"pages/OUTPUT/{record['VENDOR']}-contract.docx"
                doc.save(output_path)
            st.success("🎉 File kuitansi telah selesai dibuat, lanjutkan unggah lagi file nominatif untuk mengunduh kuitansi")
        except:
            st.warning("Unggah dulu file nominatif")   
    form2 = st.form("Proses gabung kuitansi")
    with form2:
        excelfile2 = st.file_uploader("Unggah File Nominatif")
        submit2 = st.form_submit_button("Proses Kuitansi")
        

    if submit2:
        try:
                df = pd.read_excel(excelfile2, sheet_name="Sheet1")
                np_array = df["VENDOR"].to_numpy()
                files2 = list("pages/OUTPUT/" + (np_array) + "-contract.docx")
                #files2 = files
                #st.write(files2)
                composed = f"pages/gabung.docx"
                result = Document(files2[0])
                result.add_page_break()
                composer = Composer(result)
                for i in range(1, len(files2)):
                    doc2 = Document(files2[i])
                    if i != len(files2) -1:
                        doc2.add_page_break()
                    composer.append(doc2)
                composer.save(composed)
                with open(composed, "rb") as file:
                    st.success("🎉 File kuitansi telah selesai dibuat")
                    st.download_button(
                        label = "⬇️ Download File",
                        data=file,
                        file_name="kuitansi.docx",
                        mime="application/octet-stream",
                        key="10000009"
                    )
        except:
            st.warning("Unggah file dulu")
if option == 'Perjadin':
    form = st.form("Upload file")
    with form:
        st.markdown(
        "**PERHATIAN: ** Untuk menggunakan aplikasi ini, silahkan unduh format nominatif[ di sini](https://docs.google.com/spreadsheets/d/10LDb3d-pbSRKuio8zY-HIS2kmP03ehhn/edit#gid=1302452135). Dimohon untuk mengisi nominatif sesuai dengan format file unduhan, tanpa nama kolom dan nama sheet. "
        )
        st.write("**Ini adalah format kuitansi perjalanan dinas ke daerah, jika ingin membuat kuitansi jenis lain silahkan klik menu**")

        #st.header("Proses pembuatan kuitansi")
        excelfile = st.file_uploader("Unggah File Nominatif")
        submit = st.form_submit_button("Proses Nominatif")
    if submit:
        try:
            # Convert Excel sheet to pandas dataframe
            df = pd.read_excel(excelfile, sheet_name="Sheet1")

            # Keep only date part YYYY-MM-DD (not the time)
            df["TODAY"] = pd.to_datetime(df["TODAY"]).dt.date
            df["TODAY_IN_ONE_WEEK"] = pd.to_datetime(df["TODAY_IN_ONE_WEEK"]).dt.date
            np_array = df["VENDOR"].to_numpy()
            #st.write(list(np_array))
            # Iterate over each row in df and render word document
            for record in df.to_dict(orient="records"):
                doc = DocxTemplate(f"pages/vendor-contract.docx")
                doc.render(record)
                output_path = f"pages/OUTPUT/{record['VENDOR']}-contract.docx"
                doc.save(output_path)
            st.success("🎉 File kuitansi telah selesai dibuat, lanjutkan unggah lagi file nominatif untuk mengunduh kuitansi")
        except:
            st.warning("Unggah dulu file nominatif")   
    form2 = st.form("Proses gabung kuitansi")
    with form2:
        excelfile2 = st.file_uploader("Unggah File Nominatif")
        submit2 = st.form_submit_button("Proses Kuitansi")
        

    if submit2:
        try:
                df = pd.read_excel(excelfile2, sheet_name="Sheet1")
                np_array = df["VENDOR"].to_numpy()
                files2 = list("pages/OUTPUT/" + (np_array) + "-contract.docx")
                #files2 = files
                #st.write(files2)
                composed = f"pages/gabung.docx"
                result = Document(files2[0])
                result.add_page_break()
                composer = Composer(result)
                for i in range(1, len(files2)):
                    doc2 = Document(files2[i])
                    if i != len(files2) -1:
                        doc2.add_page_break()
                    composer.append(doc2)
                composer.save(composed)
                with open(composed, "rb") as file:
                    st.success("🎉 File kuitansi telah selesai dibuat")
                    st.download_button(
                        label = "⬇️ Download File",
                        data=file,
                        file_name="kuitansi.docx",
                        mime="application/octet-stream",
                        key="10000009"
                    )
        except:
            st.warning("Unggah file dulu")



#st.sidebar.title("Laporan Keuangan")
#lk = st.sidebar.selectbox("Buat laporan", ("","Laporan keuangan",))
#st.write('You selected:', lk)